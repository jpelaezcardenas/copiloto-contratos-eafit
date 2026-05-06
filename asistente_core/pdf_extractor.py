"""Extracción de texto de documentos PDF, Word (.docx) y Google Docs.

Usa PyMuPDF para PDFs, python-docx para Word, y la API de exportación
pública de Google para Google Docs. Incluye sanitización anti prompt-injection.
"""

import fitz          # PyMuPDF
import pdfplumber
import re
import io
import requests
from typing import Optional


# ── Google Docs ───────────────────────────────────────────────────────────────

def _extract_gdoc_id(url: str) -> Optional[str]:
    """Extrae el ID del documento de una URL de Google Docs."""
    # Formato: https://docs.google.com/document/d/{ID}/edit...
    match = re.search(r'/document/d/([a-zA-Z0-9_-]+)', url)
    return match.group(1) if match else None


def extract_text_from_google_doc(url: str) -> dict:
    """Descarga y extrae texto de un Google Doc público.

    Requiere que el documento esté compartido como 'Cualquier persona con el enlace'.
    """
    result = {"text": "", "pages": [], "tables": [], "metadata": {}, "page_count": 0, "source": "google_docs"}

    doc_id = _extract_gdoc_id(url)
    if not doc_id:
        raise ValueError("URL de Google Docs no válida. Asegúrate de copiar el enlace completo.")

    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    try:
        resp = requests.get(export_url, timeout=30)
        resp.raise_for_status()
        text = resp.text
    except requests.exceptions.HTTPError as e:
        if resp.status_code == 403:
            raise PermissionError(
                "El documento de Google no es público. "
                "Compártelo con 'Cualquier persona con el enlace puede ver' e intenta de nuevo."
            )
        raise Exception(f"Error al descargar el Google Doc: {e}")
    except Exception as e:
        raise Exception(f"No se pudo acceder al Google Doc: {e}")

    result["text"] = text
    result["pages"] = [text]
    result["page_count"] = 1
    result["metadata"] = {"source_url": url, "doc_id": doc_id}
    return result


# ── Word (.docx) ──────────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes) -> dict:
    """Extrae texto de un archivo Word (.docx)."""
    result = {"text": "", "pages": [], "tables": [], "metadata": {}, "page_count": 0, "source": "docx"}

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx no está instalado. Agrega 'python-docx' a requirements.txt.")

    if hasattr(file_bytes, 'read'):
        data = file_bytes.read()
    else:
        data = file_bytes

    doc = Document(io.BytesIO(data))

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Extraer tablas
    tables_data = []
    for t in doc.tables:
        table_rows = []
        for row in t.rows:
            table_rows.append([cell.text for cell in row.cells])
        tables_data.append({"page": 1, "data": table_rows})

    full_text = "\n".join(paragraphs)
    result["text"] = full_text
    result["pages"] = [full_text]
    result["tables"] = tables_data
    result["page_count"] = 1
    result["metadata"] = {}
    return result


# ── PDF ───────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_file) -> dict:
    """Extrae texto y tablas de un PDF.

    Args:
        pdf_file: Objeto de archivo de Streamlit o bytes.

    Returns:
        dict con 'text', 'pages', 'tables', 'metadata', 'page_count'.
    """
    result = {"text": "", "pages": [], "tables": [], "metadata": {}, "page_count": 0, "source": "pdf"}

    if hasattr(pdf_file, 'read'):
        pdf_bytes = pdf_file.read()
    else:
        pdf_bytes = pdf_file

    # Extracción con PyMuPDF
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    result["page_count"] = len(doc)
    result["metadata"] = {
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "creation_date": doc.metadata.get("creationDate", ""),
    }

    page_texts = [page.get_text("text") for page in doc]
    doc.close()

    result["pages"] = page_texts
    result["text"] = "\n\n".join(page_texts)

    # Extracción de tablas con pdfplumber
    try:
        with pdfplumber.open(stream=io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                for table in (page.extract_tables() or []):
                    if table:
                        result["tables"].append({"page": i + 1, "data": table})
    except Exception:
        pass

    return result


# ── Router universal ──────────────────────────────────────────────────────────

def extract_text_from_file(uploaded_file) -> dict:
    """Router: detecta el tipo de archivo y extrae el texto apropiadamente."""
    name = uploaded_file.name.lower()
    if name.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    elif name.endswith('.docx'):
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError(f"Formato no soportado: {name}. Usa PDF o DOCX.")


# ── Sanitización ──────────────────────────────────────────────────────────────

def sanitize_extracted_text(text: str) -> str:
    """Limpia texto extraído para prevenir prompt injection."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    text = re.sub(r"[\u200b\u200c\u200d\u200e\u200f\ufeff]", "", text)

    injection_patterns = [
        r"(?i)ignor[ae]\s+(todas?\s+)?las?\s+instrucciones?",
        r"(?i)ignore\s+(all\s+)?(previous\s+)?instructions?",
        r"(?i)olvida\s+todo\s+lo\s+anterior",
        r"(?i)forget\s+(all\s+)?previous",
        r"(?i)system\s*prompt",
        r"(?i)new\s+instructions?:",
        r"(?i)override\s+mode",
    ]
    for pattern in injection_patterns:
        text = re.sub(pattern, "[CONTENIDO SOSPECHOSO ELIMINADO]", text)

    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r" {3,}", "  ", text)
    return text.strip()


def get_document_summary(extraction: dict) -> str:
    """Genera un resumen rápido del documento extraído."""
    text = extraction["text"]
    word_count = len(text.split())
    table_count = len(extraction["tables"])
    source = extraction.get("source", "documento")
    label = {"pdf": "PDF", "docx": "Word (.docx)", "google_docs": "Google Docs"}.get(source, source)

    return (
        f"**Documento procesado ({label})**\n"
        f"- Palabras: {word_count:,}\n"
        f"- Tablas encontradas: {table_count}\n"
    )
